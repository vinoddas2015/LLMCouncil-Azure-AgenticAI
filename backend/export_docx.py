"""DOCX export — generates a formatted Word document from a conversation.

Optionally adds AI hero imagery via Google Imagen 3 when GOOGLE_API_KEY is set.
Falls back to text-only if the key is missing or the call fails.
"""

import base64
import io
import re
from typing import Optional, Dict

import httpx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from .config import GOOGLE_API_KEY


def _strip_markdown(text: str) -> str:
    """Lightly strip markdown formatting for plain-text paragraphs."""
    # Bold / italic
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Links [text](url) → text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Images
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'[Image: \1]', text)
    return text


def _add_markdown_runs(paragraph, text: str):
    """Parse bold/italic in *text* and add formatted runs to *paragraph*."""
    # Split on bold markers **...**
    parts = re.split(r'(\*\*[^*]+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for italic *...*
            sub_parts = re.split(r'(\*[^*]+?\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*') and len(sp) > 2:
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sp)


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


BAYER_GREEN = RGBColor(0x10, 0x85, 0x7F)
BAYER_DARK = RGBColor(0x1E, 0x29, 0x3B)
STAGE_COLORS = {
    1: 'E8F4FD',   # light blue
    2: 'FFF3E0',   # light amber
    3: 'E8F5E9',   # light green
}


def _generate_stage_image(prompt: str) -> Optional[bytes]:
    """Call Google Imagen 3 to generate a base64 image. Returns PNG bytes or None."""
    if not GOOGLE_API_KEY:
        return None

    url = "https://generativelanguage.googleapis.com/v1beta/models/imagen-3.0:generateImage"
    headers = {"x-goog-api-key": GOOGLE_API_KEY, "Content-Type": "application/json"}
    payload = {
        "prompt": {"text": prompt[:500]},
        "aspectRatio": "16:9",
        "negativePrompt": "watermark, text, logo, words, artifacts",
    }

    try:
        with httpx.Client(timeout=40, verify=False) as client:
            resp = client.post(url, headers=headers, json=payload)
            resp.raise_for_status()
            data = resp.json()
            images = data.get("images") or data.get("generatedImages") or []
            if not images:
                return None
            b64 = images[0].get("image") or images[0].get("base64Data")
            if not b64:
                return None
            return base64.b64decode(b64)
    except Exception:
        return None


def generate_docx(conversation: dict) -> bytes:
    """Build a DOCX file from a conversation dict and return raw bytes."""
    doc = Document()

    # ── Page setup (A4) ──────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Styles ───────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = BAYER_DARK

    # ── Title ────────────────────────────────────────────────────
    title_para = doc.add_heading(conversation.get('title', 'LLM Council Conversation'), level=0)
    for run in title_para.runs:
        run.font.color.rgb = BAYER_GREEN

    # Meta
    created = conversation.get('created_at', '')
    if created:
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = meta.add_run(f'Created: {created}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')  # spacer

    # ── Optional hero image (derived from first user message) ───
    user_question = ""
    for msg in conversation.get('messages', []):
        if msg.get('role') == 'user' and msg.get('content'):
            user_question = msg.get('content', '')
            break

    hero_img = _generate_stage_image(f"Executive summary illustration for: {user_question}") if user_question else None
    if hero_img:
        try:
            doc.add_picture(io.BytesIO(hero_img), width=Cm(15))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('')
        except Exception:
            pass

    # ── Messages ─────────────────────────────────────────────────
    for msg in conversation.get('messages', []):
        if msg.get('role') == 'user':
            # User message
            h = doc.add_heading('User', level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            for line in msg.get('content', '').split('\n'):
                p = doc.add_paragraph()
                _add_markdown_runs(p, line)

        else:
            # Assistant (council) response
            h = doc.add_heading('Council Response', level=1)
            for run in h.runs:
                run.font.color.rgb = BAYER_GREEN

            # ── Stage 1 ─────────────────────────────────────────
            stage1 = msg.get('stage1')
            if stage1:
                sh = doc.add_heading('Stage 1 — Individual Model Responses', level=2)
                for run in sh.runs:
                    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

                stage_img = _generate_stage_image(f"Model responses, brainstorming, pharma context for: {user_question}") if user_question else None
                if stage_img:
                    try:
                        doc.add_picture(io.BytesIO(stage_img), width=Cm(15))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

                for resp in stage1:
                    model_name = resp.get('model', 'Unknown')
                    # Model name as bold paragraph
                    mp = doc.add_paragraph()
                    mr = mp.add_run(model_name)
                    mr.bold = True
                    mr.font.size = Pt(11)
                    mr.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

                    text = resp.get('response', 'No response')
                    for line in text.split('\n'):
                        line = line.strip()
                        if not line:
                            continue
                        # Detect markdown headings
                        if line.startswith('###'):
                            p = doc.add_heading(line.lstrip('#').strip(), level=4)
                        elif line.startswith('##'):
                            p = doc.add_heading(line.lstrip('#').strip(), level=3)
                        elif line.startswith('#'):
                            p = doc.add_heading(line.lstrip('#').strip(), level=3)
                        elif line.startswith('- ') or line.startswith('* '):
                            p = doc.add_paragraph(style='List Bullet')
                            _add_markdown_runs(p, line[2:])
                        elif re.match(r'^\d+[\.\)]\s', line):
                            p = doc.add_paragraph(style='List Number')
                            _add_markdown_runs(p, re.sub(r'^\d+[\.\)]\s', '', line))
                        else:
                            p = doc.add_paragraph()
                            _add_markdown_runs(p, line)

                    doc.add_paragraph('')  # spacer between models

            # ── Stage 2 ─────────────────────────────────────────
            stage2 = msg.get('stage2')
            if stage2:
                sh = doc.add_heading('Stage 2 — Peer Rankings', level=2)
                for run in sh.runs:
                    run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

                stage_img = _generate_stage_image(f"Peer review, ranking board, evaluation visuals for: {user_question}") if user_question else None
                if stage_img:
                    try:
                        doc.add_picture(io.BytesIO(stage_img), width=Cm(15))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

                for ranking in stage2:
                    model_name = ranking.get('model', 'Unknown')
                    mp = doc.add_paragraph()
                    mr = mp.add_run(model_name)
                    mr.bold = True
                    mr.font.color.rgb = RGBColor(0xBF, 0x36, 0x0C)

                    text = ranking.get('response', 'No ranking')
                    for line in text.split('\n'):
                        line = line.strip()
                        if not line:
                            continue
                        if line.startswith('- ') or line.startswith('* '):
                            p = doc.add_paragraph(style='List Bullet')
                            _add_markdown_runs(p, line[2:])
                        elif re.match(r'^\d+[\.\)]\s', line):
                            p = doc.add_paragraph(style='List Number')
                            _add_markdown_runs(p, re.sub(r'^\d+[\.\)]\s', '', line))
                        else:
                            p = doc.add_paragraph()
                            _add_markdown_runs(p, line)

                    doc.add_paragraph('')

            # ── Stage 3 ─────────────────────────────────────────
            stage3 = msg.get('stage3')
            if stage3:
                sh = doc.add_heading("Stage 3 — Chairman's Final Synthesis", level=2)
                for run in sh.runs:
                    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

                stage_img = _generate_stage_image(f"Executive summary, synthesized decision, pharma context for: {user_question}") if user_question else None
                if stage_img:
                    try:
                        doc.add_picture(io.BytesIO(stage_img), width=Cm(15))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

                chairman = stage3.get('model', 'Chairman')
                mp = doc.add_paragraph()
                mr = mp.add_run(chairman)
                mr.bold = True
                mr.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

                text = stage3.get('response', 'No synthesis')
                for line in text.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith('###'):
                        p = doc.add_heading(line.lstrip('#').strip(), level=4)
                    elif line.startswith('##'):
                        p = doc.add_heading(line.lstrip('#').strip(), level=3)
                    elif line.startswith('#'):
                        p = doc.add_heading(line.lstrip('#').strip(), level=3)
                    elif line.startswith('- ') or line.startswith('* '):
                        p = doc.add_paragraph(style='List Bullet')
                        _add_markdown_runs(p, line[2:])
                    elif re.match(r'^\d+[\.\)]\s', line):
                        p = doc.add_paragraph(style='List Number')
                        _add_markdown_runs(p, re.sub(r'^\d+[\.\)]\s', '', line))
                    else:
                        p = doc.add_paragraph()
                        _add_markdown_runs(p, line)

        # Horizontal rule between messages
        doc.add_paragraph('─' * 60)

    # ── Footer ───────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('Generated by LLM Council — Bayer myGenAssist')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Serialise to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
